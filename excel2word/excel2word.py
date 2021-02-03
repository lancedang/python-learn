##pip install openpyxl
from openpyxl import load_workbook
#pip install python-docx
from docx import Document
#pip install shutil
from shutil import copyfile
import os

datadir = 'E:/pylearn/excel2word/'
excelname='data.xlsx'
templatename='template.docx'
newdocdir=datadir+'all/'



def editWordTable():
    # 实例化
    #docname = datadir+'result.docx'

    copyfile('E:/pylearn/template.docx', 'E:/pylearn/1122332.docx')

    document = Document('E:/pylearn/1122332.docx')
    # 读取word中的所有表格
    table = document.tables[0]
    print('name=' + table.cell(0,3).text)
    print('sex=' + table.cell(0,7).text)
    print('age=' + table.cell(0,12).text)
    print('id=' + table.cell(1,3).text)
    print('phone=' + table.cell(1,12).text)
    print('address=' + table.cell(2,3).text)
    table.cell(2,3).text = '111111'
    document.save('E:/pylearn/1122332.docx')

def copyexcel2word():
    # 隔离人员数据
    path = datadir + excelname  # 形成文件夹的路径便后续重复使用

    workbook = load_workbook(filename=path)
    for sheet in workbook.worksheets:
        # print(sheet.dimensions)
        # A1:W10
        for rowNumber in range(3, sheet.max_row + 1):
            # 获取隔离人员excel 数据
            name = sheet.cell(row=rowNumber, column=4).value
            sex = sheet.cell(row=rowNumber, column=5).value
            age = sheet.cell(row=rowNumber, column=6).value
            idno = sheet.cell(row=rowNumber, column=7).value
            phone = sheet.cell(row=rowNumber, column=8).value
            address = sheet.cell(row=rowNumber, column=9).value

            # 根据template创建/复制一个新docx文件
            templatedocfile = datadir + templatename
            # 以行号-name.docx命名
            newdocfile = newdocdir + sheet.title + '-' + str(rowNumber) + '-' + name + '.docx'
            copyfile(templatedocfile, newdocfile)

            document = Document(newdocfile)
            # 读取word中的所有表格
            table = document.tables[0]

            table.cell(0, 3).text = str(name)
            table.cell(0, 7).text = str(sex)
            table.cell(0, 12).text = str(age)
            table.cell(1, 3).text = str(idno)
            table.cell(1, 12).text = str(phone)
            table.cell(2, 3).text = str(address)

            # 保存新数据
            document.save(newdocfile)

            """
            print('name=' + table.cell(0, 3).text)
            print('sex=' + table.cell(0, 7).text)
            print('age=' + table.cell(0, 12).text)
            print('id=' + table.cell(1, 3).text)
            print('phone=' + table.cell(1, 12).text)
            print('address=' + table.cell(2, 3).text)
            """

            # table.cell(0, 3).text = '111111'
            # document.save(docname)


if __name__ == '__main__':
    # editWordTable()
    copyexcel2word()

